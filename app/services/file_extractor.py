"""
File extraction service supporting PDF, DOCX, TXT, and image formats.
Handles diverse resume formats with fallback strategies.
Image-based resumes (PNG, JPG, WEBP) are processed via Tesseract OCR.
"""

import io
import re
from pathlib import Path
from typing import Optional, Tuple

import chardet
import pdfplumber
from PyPDF2 import PdfReader
from docx import Document
from loguru import logger

from app.core.config import settings
from app.core.exceptions import (
    FileExtractionError,
    UnsupportedFileFormatError,
    FileSizeLimitError,
    EmptyResumeError,
)

# Lazy-loaded OCR dependencies
_pytesseract = None
_PIL_Image = None
_ImageEnhance = None
_ImageFilter = None
_ocr_available: Optional[bool] = None


def _load_ocr_deps() -> bool:
    """Lazy-load OCR dependencies and configure Tesseract path."""
    global _pytesseract, _PIL_Image, _ImageEnhance, _ImageFilter, _ocr_available

    if _ocr_available is not None:
        return _ocr_available

    try:
        import pytesseract
        from PIL import Image, ImageEnhance, ImageFilter

        # Configure Tesseract path if specified
        if settings.TESSERACT_CMD:
            pytesseract.pytesseract.tesseract_cmd = settings.TESSERACT_CMD

        # Verify Tesseract is actually reachable
        pytesseract.get_tesseract_version()

        _pytesseract = pytesseract
        _PIL_Image = Image
        _ImageEnhance = ImageEnhance
        _ImageFilter = ImageFilter
        _ocr_available = True

        logger.info(
            f"OCR engine ready — Tesseract "
            f"v{pytesseract.get_tesseract_version()}"
        )
    except Exception as e:
        _ocr_available = False
        logger.warning(
            f"OCR unavailable — Tesseract not found: {e}. "
            "Image resume uploads will be rejected. "
            "Install Tesseract: https://github.com/tesseract-ocr/tesseract"
        )

    return _ocr_available


class FileExtractor:
    """
    Extracts raw text from resume files.
    Implements multiple extraction strategies per format with fallback logic.
    Supports image-based resumes (PNG, JPG, WEBP) via Tesseract OCR.
    """

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".doc"}
    IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp"}

    def __init__(self):
        logger.info("FileExtractor initialized")

    async def extract(
        self, file_content: bytes, filename: str
    ) -> Tuple[str, str]:
        """
        Extract text from uploaded file.

        Args:
            file_content: Raw bytes of the uploaded file
            filename: Original filename with extension

        Returns:
            Tuple of (extracted_text, file_type)

        Raises:
            UnsupportedFileFormatError: If file type not supported
            FileSizeLimitError: If file exceeds size limit
            FileExtractionError: If extraction fails
            EmptyResumeError: If no text extracted
        """
        # Validate file size
        file_size_mb = len(file_content) / (1024 * 1024)
        if file_size_mb > settings.MAX_FILE_SIZE_MB:
            raise FileSizeLimitError(
                f"File size {file_size_mb:.1f}MB exceeds limit "
                f"of {settings.MAX_FILE_SIZE_MB}MB",
                details={"file_size_mb": file_size_mb}
            )

        # Determine file extension
        ext = Path(filename).suffix.lower()
        all_supported = self.SUPPORTED_EXTENSIONS | self.IMAGE_EXTENSIONS
        if ext not in all_supported:
            raise UnsupportedFileFormatError(
                f"Unsupported file format: '{ext}'. "
                f"Supported: {all_supported}",
                details={"extension": ext}
            )

        logger.info(
            f"Extracting text from '{filename}' "
            f"({file_size_mb:.2f} MB, type: {ext})"
        )

        # Route to appropriate extractor
        try:
            if ext == ".pdf":
                text = await self._extract_pdf(file_content)
            elif ext == ".docx":
                text = await self._extract_docx(file_content)
            elif ext in (".txt", ".doc"):
                text = await self._extract_text(file_content)
            elif ext in self.IMAGE_EXTENSIONS:
                text = await self._extract_image(file_content, filename)
            else:
                raise UnsupportedFileFormatError(
                    f"No extractor for: {ext}"
                )
        except (UnsupportedFileFormatError, FileSizeLimitError):
            raise
        except Exception as e:
            logger.error(f"Extraction failed for {filename}: {e}")
            raise FileExtractionError(
                f"Failed to extract text from '{filename}': {str(e)}",
                details={"filename": filename, "error": str(e)}
            )

        # Validate extraction result
        cleaned_text = self._clean_extracted_text(text)
        if not cleaned_text or len(cleaned_text.strip()) < 50:
            raise EmptyResumeError(
                "Resume appears empty or contains insufficient text. "
                f"Extracted only {len(cleaned_text)} characters.",
                details={
                    "extracted_length": len(cleaned_text),
                    "filename": filename
                }
            )

        logger.info(
            f"Successfully extracted {len(cleaned_text)} characters "
            f"from '{filename}'"
        )
        return cleaned_text, ext

    async def _extract_image(self, content: bytes, filename: str) -> str:
        """
        Extract text from image-based resumes using Tesseract OCR.
        Applies pre-processing (grayscale, contrast, sharpening) to
        handle decorative Canva-style backgrounds gracefully.
        """
        if not _load_ocr_deps():
            raise FileExtractionError(
                "Cannot process image resumes — Tesseract OCR is not "
                "installed. Please install Tesseract and restart the server. "
                "Windows: choco install tesseract | "
                "Linux: apt-get install tesseract-ocr | "
                "macOS: brew install tesseract",
                details={"filename": filename}
            )

        try:
            # Open image with Pillow
            image = _PIL_Image.open(io.BytesIO(content))
            logger.info(
                f"OCR: Opened image {image.size[0]}x{image.size[1]} "
                f"mode={image.mode} format={image.format}"
            )

            # Pre-process for better OCR accuracy on Canva resumes
            processed = self._preprocess_image_for_ocr(image)

            # Run Tesseract OCR
            ocr_config = (
                f"--dpi {settings.OCR_DPI} "
                f"--oem 3 --psm 6"
            )
            text = _pytesseract.image_to_string(
                processed,
                lang=settings.OCR_LANGUAGE,
                config=ocr_config,
            )

            char_count = len(text.strip()) if text else 0
            logger.info(
                f"OCR: Extracted {char_count} characters from '{filename}'"
            )

            if not text or char_count < 10:
                # Retry with different page segmentation mode (full auto)
                logger.info("OCR: Low yield, retrying with PSM 3 (full auto)...")
                ocr_config_retry = (
                    f"--dpi {settings.OCR_DPI} "
                    f"--oem 3 --psm 3"
                )
                text = _pytesseract.image_to_string(
                    processed,
                    lang=settings.OCR_LANGUAGE,
                    config=ocr_config_retry,
                )

            return text or ""

        except Exception as e:
            logger.error(f"OCR extraction failed for '{filename}': {e}")
            raise FileExtractionError(
                f"OCR extraction failed for '{filename}': {str(e)}",
                details={"filename": filename, "error": str(e)}
            )

    def _preprocess_image_for_ocr(self, image):
        """
        Pre-process an image for optimal OCR accuracy.
        Handles Canva-style resumes with decorative backgrounds,
        colored text, and complex layouts.
        """
        # Convert to RGB if necessary (handles RGBA, palette, etc.)
        if image.mode not in ("RGB", "L"):
            image = image.convert("RGB")

        # Resize small images to improve OCR accuracy
        width, height = image.size
        if width < 1500:
            scale = 1500 / width
            new_size = (int(width * scale), int(height * scale))
            image = image.resize(new_size, _PIL_Image.LANCZOS)
            logger.debug(
                f"OCR pre-process: upscaled to {new_size[0]}x{new_size[1]}"
            )

        # Convert to grayscale
        gray = image.convert("L")

        # Enhance contrast (helps with light-colored Canva backgrounds)
        enhancer = _ImageEnhance.Contrast(gray)
        gray = enhancer.enhance(2.0)

        # Sharpen (helps with text clarity)
        enhancer = _ImageEnhance.Sharpness(gray)
        gray = enhancer.enhance(2.0)

        # Apply slight denoise filter
        gray = gray.filter(_ImageFilter.MedianFilter(size=3))

        # Binarize: convert to pure black and white
        threshold = 140
        gray = gray.point(lambda p: 255 if p > threshold else 0, "1")

        return gray

    async def _extract_pdf(self, content: bytes) -> str:
        """
        Extract text from PDF using pdfplumber (primary) 
        with PyPDF2 fallback.
        """
        text_parts = []

        # Primary: pdfplumber (better with complex layouts)
        try:
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                    else:
                        # Try extracting tables as text
                        tables = page.extract_tables()
                        for table in tables:
                            for row in table:
                                if row:
                                    row_text = " | ".join(
                                        str(cell) for cell in row if cell
                                    )
                                    text_parts.append(row_text)

            if text_parts:
                return "\n".join(text_parts)
        except Exception as e:
            logger.warning(f"pdfplumber failed, trying PyPDF2: {e}")

        # Fallback: PyPDF2
        try:
            reader = PdfReader(io.BytesIO(content))
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

            return "\n".join(text_parts)
        except Exception as e:
            raise FileExtractionError(
                f"Both PDF extractors failed: {e}"
            )

    async def _extract_docx(self, content: bytes) -> str:
        """Extract text from DOCX including paragraphs and tables."""
        try:
            doc = Document(io.BytesIO(content))
            text_parts = []

            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip()
                        for cell in row.cells
                        if cell.text.strip()
                    )
                    if row_text:
                        text_parts.append(row_text)

            return "\n".join(text_parts)
        except Exception as e:
            raise FileExtractionError(
                f"DOCX extraction failed: {e}"
            )

    async def _extract_text(self, content: bytes) -> str:
        """Extract text from plain text files with encoding detection."""
        try:
            # Detect encoding
            detection = chardet.detect(content)
            encoding = detection.get("encoding", "utf-8")
            confidence = detection.get("confidence", 0)

            logger.debug(
                f"Detected encoding: {encoding} "
                f"(confidence: {confidence:.2f})"
            )

            # Try detected encoding first
            try:
                return content.decode(encoding)
            except (UnicodeDecodeError, TypeError):
                pass

            # Fallback encodings
            for enc in ["utf-8", "latin-1", "ascii", "cp1252"]:
                try:
                    return content.decode(enc)
                except UnicodeDecodeError:
                    continue

            # Last resort: decode with replacement
            return content.decode("utf-8", errors="replace")

        except Exception as e:
            raise FileExtractionError(
                f"Text extraction failed: {e}"
            )

    def _clean_extracted_text(self, text: str) -> str:
        """Basic cleaning of extracted text."""
        if not text:
            return ""

        # Replace multiple whitespace but preserve newlines
        text = re.sub(r'[^\S\n]+', ' ', text)
        # Remove excessive blank lines
        text = re.sub(r'\n{4,}', '\n\n\n', text)
        # Remove null bytes and control characters
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', text)
        # Strip unmapped glyph placeholders like "(cid:131)" that pdfplumber
        # emits for icon/symbol fonts in complex resume layouts.
        text = re.sub(r'\(cid:\d+\)', '', text)
        # Collapse any spaces left behind by the removals.
        text = re.sub(r'[^\S\n]{2,}', ' ', text)

        return text.strip()